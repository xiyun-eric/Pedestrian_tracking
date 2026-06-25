"""提取docx文件内容"""
from docx import Document

doc = Document('docs/动态场景多目标跟踪系统设计与实现 (2).docx')
with open('docs/report_v2_content.txt', 'w', encoding='utf-8') as f:
    for i, para in enumerate(doc.paragraphs):
        f.write(f'P{i}: {para.text}\n')
    f.write('\n=== TABLES ===\n')
    for ti, table in enumerate(doc.tables):
        f.write(f'\n--- Table {ti} ---\n')
        for ri, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            f.write(f'  Row {ri}: {" | ".join(cells)}\n')

print(f'Done, paragraphs: {len(doc.paragraphs)}, tables: {len(doc.tables)}')
